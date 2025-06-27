"""
FastAPI API routes for document list and upload.
"""
from fastapi import APIRouter, UploadFile, File, Request, HTTPException, BackgroundTasks, Body, Query
from fastapi.responses import JSONResponse, FileResponse
from fastapi.templating import Jinja2Templates
from pathlib import Path
from typing import List
import concurrent.futures
import logging
import traceback
import tempfile
import markdown
import pdfkit
from fpdf import FPDF

from app.services.openai_service import generate_funcional_analysis
from docx import Document as DocxDocument
from docx.shared import Pt

router = APIRouter()
templates = Jinja2Templates(directory="templates")
UPLOAD_DIR = Path("static/uploads")
UPLOAD_DIR.mkdir(parents=True, exist_ok=True)

documents_cache: List[str] = []

# Configuración básica de logging a fichero
logging.basicConfig(
    filename="app_error.log",
    level=logging.ERROR,
    format="%(asctime)s %(levelname)s %(message)s",
)

@router.get("/documents/list", response_class=JSONResponse)
async def get_documents_list() -> list[str]:
    """Return the list of uploaded documents, excluyendo el funcional generado."""
    files = [f.name for f in UPLOAD_DIR.iterdir() if f.is_file() and f.name != "funcional_generado.md"]
    return files

@router.post("/upload")
async def upload_files(files: List[UploadFile] = File(...)):
    """Handle file uploads."""
    saved = []
    for file in files:
        dest = UPLOAD_DIR / file.filename
        with dest.open("wb") as f:
            f.write(await file.read())
        saved.append(file.filename)
    return {"uploaded": saved}

@router.get("/documents")
def documents_page(request: Request):
    """Render the document list page."""
    return templates.TemplateResponse("documents/list.html", {"request": request})

@router.delete("/documents/{filename}", response_class=JSONResponse)
def delete_document(filename: str):
    """Delete a document by filename."""
    file_path = UPLOAD_DIR / filename
    if file_path.exists() and file_path.is_file():
        file_path.unlink()
        return {"deleted": filename}
    raise HTTPException(status_code=404, detail="Documento no encontrado")

@router.post("/documents/generate-funcional", response_class=JSONResponse)
async def generate_funcional(background_tasks: BackgroundTasks):
    """Genera un análisis funcional a partir de los documentos subidos, lo devuelve en memoria y lo guarda en un archivo interno."""
    files = [str(f) for f in UPLOAD_DIR.iterdir() if f.is_file()]
    if not files:
        return {"error": "No hay documentos para analizar."}
    loop = None
    try:
        import asyncio
        loop = asyncio.get_running_loop()
    except RuntimeError:
        pass
    def sync_generate():
        try:
            from app.services.openai_service import generate_funcional_analysis
            return generate_funcional_analysis(files)
        except Exception as e:
            tb = traceback.format_exc()
            # Loguea el error y el traceback en un fichero
            logging.error(f"Error generando el análisis funcional: {e}\nTRACEBACK:\n{tb}")
            return f"[ERROR] {str(e)}\nTRACEBACK:\n{tb}"
    if loop:
        analysis = await loop.run_in_executor(None, sync_generate)
    else:
        analysis = sync_generate()
    if analysis.startswith("[ERROR]"):
        return JSONResponse(status_code=500, content={"error": analysis})
    internal_path = UPLOAD_DIR / "funcional_generado.md"
    internal_path.write_text(analysis, encoding="utf-8")
    return {"funcional": analysis}

@router.post("/documents/update-funcional", response_class=JSONResponse)
async def update_funcional(data: dict = Body(...)) -> dict:
    """Actualiza el documento funcional en memoria y en el archivo funcional_generado.md."""
    content = data.get("content", "")
    internal_path = UPLOAD_DIR / "funcional_generado.md"
    internal_path.write_text(content, encoding="utf-8")
    return {"ok": True}

@router.get("/documents/export-funcional")
async def export_funcional(format: str = Query("docx", enum=["docx", "pdf"])):
    """Exporta el documento funcional a Word o PDF con formato interpretado desde Markdown."""
    internal_path = UPLOAD_DIR / "funcional_generado.md"
    if not internal_path.exists():
        return JSONResponse(status_code=404, content={"error": "No existe el documento funcional generado."})
    content = internal_path.read_text(encoding="utf-8")
    if format == "docx":
        import markdown
        from bs4 import BeautifulSoup
        from docx import Document as DocxDocument
        import tempfile
        html = markdown.markdown(content)
        soup = BeautifulSoup(html, "html.parser")
        doc = DocxDocument()
        # --- Insertar solo la Tabla de Contenido nativa de Word (TOC) ---
        from docx.oxml import OxmlElement
        from docx.oxml.ns import qn
        p = doc.add_paragraph()
        run = p.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.text = r'TOC \o "1-3" \h \z \u'
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'end')
        r_element = run._r
        r_element.append(fldChar1)
        r_element.append(instrText)
        r_element.append(fldChar2)
        doc.add_paragraph("")  # Espacio tras el TOC
        # --- Insertar el resto del documento ---
        for el in soup.children:
            if el.name and el.name.startswith('h') and el.name[1:].isdigit():
                level = int(el.name[1:])
                doc.add_heading(el.get_text(), level=level if level <= 4 else 4)
            elif el.name == 'ul':
                for li in el.find_all('li', recursive=False):
                    p = doc.add_paragraph(li.get_text(), style='List Bullet')
            elif el.name == 'ol':
                for li in el.find_all('li', recursive=False):
                    p = doc.add_paragraph(li.get_text(), style='List Number')
            elif el.name == 'p':
                p = doc.add_paragraph(el.get_text())
            elif el.name == 'strong':
                p = doc.add_paragraph()
                run = p.add_run(el.get_text())
                run.bold = True
            elif el.name == 'em':
                p = doc.add_paragraph()
                run = p.add_run(el.get_text())
                run.italic = True
            # Puedes agregar más reglas para otros elementos si lo necesitas
        # No se añade ningún índice textual ni contenido relacionado con el árbol de contenido al final
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            doc.save(tmp.name)
            tmp.flush()
            return FileResponse(tmp.name, filename="funcional_generado.docx", media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    elif format == "pdf":
        import tempfile
        import re
        class PDF(FPDF):
            def header(self):
                self.set_font('Arial', 'B', 14)
                self.set_text_color(30, 64, 175)
                self.cell(0, 10, 'Documento Funcional', ln=True, align='C')
                self.ln(4)
        pdf = PDF()
        pdf.add_page()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.set_font('Arial', '', 12)
        bullet = '-'  # O '*'
        lines = content.split('\n')

        MAX_CHARS_PER_LINE = 120

        def force_fit_line(line: str, pdf: FPDF, max_width: float) -> list[str]:
            """
            Divide la línea en fragmentos que nunca exceden el ancho de la celda.
            Si una palabra es demasiado larga, la trocea carácter a carácter.
            Si ni un solo carácter imprimible cabe, pone '?'.
            """
            safe_chunks = []
            current = ''
            widest = max((pdf.get_string_width(chr(c)) for c in range(32, 127)), default=1)
            if max_width < widest or max_width < 1:
                return ['?']
            for word in line.split(' '):
                while word:
                    if pdf.get_string_width(word) <= max_width:
                        if current:
                            test = current + ' ' + word
                        else:
                            test = word
                        if pdf.get_string_width(test) <= max_width:
                            current = test
                            word = ''
                        else:
                            if current:
                                safe_chunks.append(current)
                                current = ''
                            else:
                                i = 1
                                while i <= len(word):
                                    if pdf.get_string_width(word[:i]) > max_width:
                                        if i == 1:
                                            safe_chunks.append('?')
                                            word = word[1:]
                                        else:
                                            safe_chunks.append(word[:i-1])
                                            word = word[i-1:]
                                        break
                                    i += 1
                                else:
                                    safe_chunks.append(word)
                                    word = ''
                    else:
                        i = 1
                        while i <= len(word):
                            if pdf.get_string_width(word[:i]) > max_width:
                                if i == 1:
                                    safe_chunks.append('?')
                                    word = word[1:]
                                else:
                                    safe_chunks.append(word[:i-1])
                                    word = word[i-1:]
                                break
                            i += 1
                        else:
                            safe_chunks.append(word)
                            word = ''
            if current:
                safe_chunks.append(current)
            return safe_chunks

        def split_long_line(line: str, max_chars: int) -> list[str]:
            """
            Divide una línea en sublíneas de máximo max_chars caracteres.
            """
            return [line[i:i+max_chars] for i in range(0, len(line), max_chars)]

        for line in lines:
            line = line.rstrip('\r').strip()
            if not line:
                pdf.ln(4)
                continue
            # Reemplaza cualquier tabulación por espacio
            line = line.replace('\t', ' ')
            # Reemplaza cualquier carácter no imprimible por espacio
            line = ''.join(c if 32 <= ord(c) <= 126 else ' ' for c in line)
            max_width = pdf.w - pdf.l_margin - pdf.r_margin
            # Limita la longitud de la línea antes de procesar
            sublines = split_long_line(line, MAX_CHARS_PER_LINE)
            for subline in sublines:
                # Encabezados
                m = re.match(r'^(#+)\s*(.*)', subline)
                if m:
                    level = len(m.group(1))
                    text = m.group(2)
                    if level == 1:
                        pdf.set_font('Arial', 'B', 16)
                        pdf.set_text_color(30, 64, 175)
                        for chunk in force_fit_line(text, pdf, max_width):
                            if not chunk.strip() or pdf.get_string_width(chunk) > max_width:
                                chunk = '?'
                            pdf.multi_cell(0, 10, chunk)
                        pdf.set_font('Arial', '', 12)
                        pdf.set_text_color(34, 34, 34)
                    elif level == 2:
                        pdf.set_font('Arial', 'B', 14)
                        pdf.set_text_color(30, 64, 175)
                        for chunk in force_fit_line(text, pdf, max_width):
                            if not chunk.strip() or pdf.get_string_width(chunk) > max_width:
                                chunk = '?'
                            pdf.multi_cell(0, 8, chunk)
                        pdf.set_font('Arial', '', 12)
                        pdf.set_text_color(34, 34, 34)
                    elif level == 3:
                        pdf.set_font('Arial', 'B', 12)
                        pdf.set_text_color(30, 64, 175)
                        for chunk in force_fit_line(text, pdf, max_width):
                            if not chunk.strip() or pdf.get_string_width(chunk) > max_width:
                                chunk = '?'
                            pdf.multi_cell(0, 7, chunk)
                        pdf.set_font('Arial', '', 12)
                        pdf.set_text_color(34, 34, 34)
                    else:
                        pdf.set_font('Arial', 'B', 11)
                        pdf.set_text_color(30, 64, 175)
                        for chunk in force_fit_line(text, pdf, max_width):
                            if not chunk.strip() or pdf.get_string_width(chunk) > max_width:
                                chunk = '?'
                            pdf.multi_cell(0, 6, chunk)
                        pdf.set_font('Arial', '', 12)
                        pdf.set_text_color(34, 34, 34)
                    continue
                # Listas
                if subline.startswith('- '):
                    pdf.set_x(pdf.get_x() + 5)
                    for chunk in force_fit_line(f'{bullet} ' + subline[2:], pdf, max_width - 5):
                        if not chunk.strip() or pdf.get_string_width(chunk) > max_width - 5:
                            chunk = '?'
                        pdf.multi_cell(0, 7, chunk)
                    continue
                # Negrita
                bold = re.match(r'^\*\*(.+)\*\*$', subline)
                if bold:
                    pdf.set_font('Arial', 'B', 12)
                    for chunk in force_fit_line(bold.group(1), pdf, max_width):
                        if not chunk.strip() or pdf.get_string_width(chunk) > max_width:
                            chunk = '?'
                        pdf.multi_cell(0, 7, chunk)
                    pdf.set_font('Arial', '', 12)
                    continue
                # Cursiva
                italic = re.match(r'^\*(.+)\*$', subline)
                if italic:
                    pdf.set_font('Arial', 'I', 12)
                    for chunk in force_fit_line(italic.group(1), pdf, max_width):
                        if not chunk.strip() or pdf.get_string_width(chunk) > max_width:
                            chunk = '?'
                        pdf.multi_cell(0, 7, chunk)
                    pdf.set_font('Arial', '', 12)
                    continue
                # Texto normal
                for chunk in force_fit_line(subline, pdf, max_width):
                    if not chunk.strip() or pdf.get_string_width(chunk) > max_width:
                        chunk = '?'
                    pdf.multi_cell(0, 7, chunk)
    return JSONResponse(status_code=400, content={"error": "Formato no soportado."})
